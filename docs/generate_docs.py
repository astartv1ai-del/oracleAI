"""Пересборка «Оракул_документация.docx/.pdf» из markdown-источников docs/.

Использование:
    .venv/bin/python docs/generate_docs.py            # docx + pdf
    .venv/bin/python docs/generate_docs.py --pdf      # только pdf
    .venv/bin/python docs/generate_docs.py --docx     # только docx

Зависимости (venv): python-docx, markdown, weasyprint.
Порядок разделов = оглавление сборного документа (README первым).
"""
from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

import markdown
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

ROOT = Path(__file__).resolve().parent.parent
DOCS = Path(__file__).resolve().parent
OUT_NAME = "Оракул_документация"

#: Разделы сборного документа и их файлы-источники.
SECTIONS = [
    ("Обзор (README)", "README.md"),
    ("Идея продукта", "IDEA.md"),
    ("Продукт и функциональность", "PRODUCT.md"),
    ("Маркетинг", "MARKETING.md"),
    ("Разработка", "DEVELOPMENT.md"),
    ("Дизайн и бренд", "DESIGN_SPEC.md"),
    ("Frontend-ТЗ", "FRONTEND_TZ.md"),
    ("Бизнес и монетизация", "MONETIZATION.md"),
    ("Операции и эксплуатация", "OPERATIONS.md"),
    ("Право и приватность", "LEGAL.md"),
]

TITLE = "Оракул — личный AI-астролог в Telegram"
SUBTITLE = ("Документация проекта: идея, продукт, маркетинг, разработка, "
            "дизайн, монетизация, операции, право")


def _load(source: str) -> str:
    if source == "README.md":
        return (ROOT / "README.md").read_text(encoding="utf-8")
    return (DOCS / source).read_text(encoding="utf-8")


# ─────────────────────────────── docx ────────────────────────────────────────

_H1 = re.compile(r"^(#{1,3})\s+(.*)$", re.M)
_TABLE = re.compile(r"^\|.*\|$", re.M)
_LIST = re.compile(r"^[-*]\s+(.*)$", re.M)
_CODE_BLOCK = re.compile(r"^```.*$", re.M)


def _md_to_docx_paragraphs(doc: Document, md: str) -> None:
    """Минимальный md→docx: заголовки, списки, код-блоки, абзацы, таблицы-текстом."""
    in_code = False
    for raw in md.splitlines():
        line = raw.rstrip()
        if _CODE_BLOCK.match(line):
            in_code = not in_code
            continue
        if in_code:
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            continue
        if not line.strip():
            continue
        m = _H1.match(line)
        if m:
            level = len(m.group(1))
            h = doc.add_heading(level=min(level, 4))
            h.add_run(_strip_md(m.group(2)))
            continue
        m = _LIST.match(line)
        if m:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(_strip_md(m.group(1)))
            continue
        if _TABLE.match(line):
            # таблицы markdown — текстом: строка как колонки через «|»
            cells = [c.strip() for c in line.strip("|").split("|")]
            if any(re.fullmatch(r":?-{2,}:?", c) for c in cells):
                continue  # разделитель заголовка
            doc.add_paragraph(" | ".join(cells), style="No Spacing")
            continue
        p = doc.add_paragraph()
        p.add_run(_strip_md(line))


def _strip_md(text: str) -> str:
    text = re.sub(r"`([^`]*)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    text = re.sub(r"\[([^\]]*)\]\([^)]*\)", r"\1", text)
    text = re.sub(r"!\[([^\]]*)\]\([^)]*\)", r"\1", text)
    return text


def build_docx() -> Path:
    doc = Document()
    # титульная страница
    t = doc.add_heading(level=0)
    r = t.add_run(TITLE)
    r.font.size = Pt(24)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(SUBTITLE).italic = True
    doc.add_page_break()
    # оглавление
    doc.add_heading("Содержание", level=1)
    for title, _ in SECTIONS:
        doc.add_paragraph(title, style="List Number")
    doc.add_page_break()

    for title, source in SECTIONS:
        doc.add_heading(title, level=1)
        _md_to_docx_paragraphs(doc, _load(source))
        doc.add_page_break()

    out = ROOT / f"{OUT_NAME}.docx"
    doc.save(out)
    return out


# ─────────────────────────────── pdf ─────────────────────────────────────────

def build_pdf() -> Path:
    body: list[str] = []
    for title, source in SECTIONS:
        md = markdown.markdown(_load(source), extensions=["tables", "fenced_code"])
        body.append(f"<h1 class=\"section\">{title}</h1>\n{md}")
    html = f"""<!doctype html><html lang="ru"><head><meta charset="utf-8">
<style>
@page {{ size: A4; margin: 1.6cm 1.7cm; @bottom-center {{
  content: counter(page) " / " counter(pages); font-size: 9px; color: #888; }} }}
body {{ font-family: 'DejaVu Sans', sans-serif; font-size: 10.5pt; line-height: 1.5;
  color: #1a1a1a; }}
h1 {{ font-size: 20pt; color: #7a5c00; border-bottom: 2px solid #d4af37;
  padding-bottom: 4px; margin: 0 0 14px; }}
h1.section {{ page-break-before: always; }}
h2 {{ font-size: 15pt; color: #4a3a66; margin-top: 18px; }}
h3 {{ font-size: 12.5pt; color: #333; }}
h4 {{ font-size: 11pt; color: #555; }}
code {{ background: #f4f1e8; padding: 1px 4px; border-radius: 3px;
  font-size: 9pt; }}
pre {{ background: #f7f5f0; border: 1px solid #e0dccf; border-radius: 4px;
  padding: 8px; font-size: 8.8pt; white-space: pre-wrap; }}
table {{ border-collapse: collapse; width: 100%; margin: 8px 0; font-size: 9.5pt; }}
th, td {{ border: 1px solid #d8d2c2; padding: 4px 7px; text-align: left; }}
th {{ background: #f4f1e8; }}
blockquote {{ border-left: 3px solid #d4af37; margin: 6px 0; padding: 2px 12px;
  color: #555; background: #faf8f2; }}
hr {{ border: none; border-top: 1px solid #ddd; }}
</style></head><body>
<h1>{TITLE}</h1>
<p><em>{SUBTITLE}</em></p>
<p><strong>Содержание:</strong><br>{ "<br>".join(
    f"{i+1}. {t}" for i, (t, _) in enumerate(SECTIONS)) }</p>
{''.join(body)}
</body></html>"""
    out = ROOT / f"{OUT_NAME}.pdf"
    from weasyprint import HTML
    HTML(string=html).write_pdf(out)
    return out


def main() -> None:
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("--pdf", action="store_true")
    ap.add_argument("--docx", action="store_true")
    args = ap.parse_args()
    if args.pdf or not args.docx:
        p = build_pdf()
        print(f"PDF:  {p} ({p.stat().st_size // 1024} КБ)")
    if args.docx or not args.pdf:
        p = build_docx()
        print(f"DOCX: {p} ({p.stat().st_size // 1024} КБ)")


if __name__ == "__main__":
    sys.exit(main())
